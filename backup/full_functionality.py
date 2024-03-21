import PyPDF2
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openai
from openai import OpenAI
import boto3

from copy import deepcopy
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import BytesIO
import os
import base64
import requests
import time
import json

api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

s3_client = boto3.client("s3")
BUCKET_NAME = os.environ.get("resume_bucket")

# open the file: "/Users/campbmaso/Desktop/Development/GitHub/Sandbox/Backend/resumes/TEMPLATE - Pscyence Resume Service (1).docx"

# =========================== get resume files ready =============================================================================================
# local version
# file_path = "/Users/campbmaso/Desktop/Development/GitHub/Sandbox/Backend/resumes/TEMPLATE - Pscyence Resume Service (1).docx"
# with open(file_path, "rb") as f:
#     file_content = BytesIO(f.read())
# new_doc = Document(file_content)
# styles = new_doc.styles

# grabbing the template resume from s3
template_version = 1
template_object = s3_client.get_object(
    Bucket="resume-s3bucket", Key=f"templates/Serif Template {template_version}.docx"
)
template_content = BytesIO(template_object["Body"].read())

new_doc = Document(template_content)

resume_text = ""
with open(
    "/Users/campbmaso/Desktop/Development/GitHub/Sandbox/Backend/resumes/Elias TOUIL CV (1).pdf",
    "rb",
) as file:
    reader = PyPDF2.PdfReader(file)
    for page in reader.pages:
        resume_text += page.extract_text()


# =========================== data templates =============================================================================================
dummy_resume = """
1
Cory G. Mazure49225 Equestrian Court, Chesterfield, MI 48047
Cell: (586) 961-5274 cory.mazure@gmail.com
LinkedIn: www.linkedin.com/in/cory-mazure
Work Experience
Process Engineer, Golf Ball R&D
Callaway Golf, Carlsbad, CA (January 2023 - Present)
● Lead engineer behind the experimentation, validation, and implementation of new spray gun equipment
in production for 3 paint lines, resulting in a more controllable process and therefore quality product
● Organized DOEs to study differences in mechanical processes between R&D and production, reducing
variance in processes, ultimately establishing confidence which accelerated new ball development
● Implemented new test methods to validate chemistry, leading to new proprietary golf ball features
● Utilized injection molding modeling software to modify tooling, improving concentricity and yield
Mechanical Engineering, Vehicle Safety Engineering
Stellantis, Auburn Hills, MI (Summer of 2019 – Summer 2022)
● Led a design for six sigma project for optimizing the amount of test dummies in use, saving $250K/year
● Reduced vehicle validation timing by developing test fixtures for high inertia door latch testing
● Analyzed channel data for 80+ test rib deflection study, resulting in global program airbag innovations
● Coordinated JD Power quality study and benchmarking for heated steering wheel (HSW) optimization
● Strengthened best practices for HSW calibration ranges, directly applied within 2 new flagship
programs, resulting in enhanced customer satisfaction
● Investigated 2 competing suppliers for hands on detection mats via Minitab statistical analysis, leading
to more competitive pricing
● Established statistical analysis of the dummy labs 4 certification machines, reducing certification timing
and associated costs by $20K/year
● Validated 2 energy absorbing knee brackets and a driver airbag, achieving higher safety ratings
Leadership Experience
● Kettering Student Government Operations Council – Director: Planned and executed over 15 events a
term, communicating progress of the council weekly to other government sub-sections
● Sigma Chi Distant Leadership Course: Learned practical strategies for remote team communication
● Golf Club – President: Ran weekly meetings with passion, and liaised contracts with local driving range
● Mindfulness Club – President: Facilitated awareness of mental health practices and built a community
● Inter Fraternity Council – Philanthropy Chair: Orchestrated 2 events to improve the image of campus
Skills
● Stellantis Sponsored Courses: Attended 5 in depth training courses for DFMEA, DFSS, and GD&T
● Moldflow injection molding simulation software
● Proficient in NX, Fusion 360, & SolidWorks: CAE and FEA
● Experienced in Minitab statistical analysis software for interpreting DOE results
● Utilized Altair Hyperworks to analyze vehicle crash data
● Avid user of Microsoft Office and Google suite for statistical analysis and project management
Education
Bachelor of Science in Mechanical Engineering
Kettering University Flint, MI (Oct. 2019 - Dec. 2022)
● Kettering University Honors Program: Completing 2 advanced projects per academic term
● Thesis Project: WorldSID ATD Rib Deflection and Seating Analysis Assessment
● GPA: 3.95/4.00, Summa Cum Laude, Dean's List every academic term
"""

json_template = {
    "personal_information": {
        "name": "John Doe",
        "email": "john.doe@example.com",
        "phone": "123-456-7890",
        "linkedin": "linkedin.com/in/johndoe",
    },
    "summary": "Experienced software engineer with a strong background in developing scalable applications and improving software development processes.",
    "skills": [
        "Python",
        "Java",
        "Docker",
        "Kubernetes",
        "React",
        "Agile methodologies",
    ],
    "additional_activities": [
        {
            "title": "Hackathon Participant",
            "bullets": [
                "Participated in a 24-hour hackathon, developing a web application that won 2nd place."
            ],
            "date": "March 2020",
        },
    ],
    "work_experience": [
        {
            "job_title": "Senior Software Engineer",
            "company": "Tech Solutions Inc.",
            "location": "San Francisco, CA",
            "start_date": "June 2018",
            "end_date": "Present",
            "bullets": [
                "Led the development of a new feature that increased user engagement by 20%",
                "Implemented a CI/CD pipeline that reduced deployment time by 50%",
                "Mentored junior developers, improving team productivity and knowledge sharing",
            ],
        },
        {
            "job_title": "Software Engineer",
            "company": "Innovate Startup",
            "location": "San Francisco, CA",
            "date_range": "June 2016 - June 2018",
            "bullets": [
                "Developed a high-traffic web application using React and Node.js",
                "Optimized database queries, reducing load times by 30%",
                "Collaborated with cross-functional teams to define, design, and ship new features",
            ],
        },
    ],
    "education": [
        {
            "degree": "Bachelor of Science in Computer Science",
            "institution": "University of Example",
            "location": "Example City",
            "date_range": "September 2010 - June 2014",
            "bullets": "Graduated with honors",
        }
    ],
    "certifications": [
        {
            "title": "Certified Kubernetes Administrator",
            "issuer": "The Linux Foundation",
            "date_issued": "July 2019",
        }
    ],
    "projects": [
        {
            "title": "Personal Portfolio Website",
            "bullets": [
                "A personal website to showcase my projects and resume. Built with HTML, CSS, and JavaScript."
            ],
            "link": "http://www.johndoeportfolio.com",
        }
    ],
}

header_template = {
    "personal_information": {
        "name": "John Doe",
        "email": "john.doe@example.com",
        "phone": "123-456-7890",
        "linkedin": "linkedin.com/in/johndoe",
    }
}
summary_template = {
    "summary": "Experienced software engineer with a strong background in developing scalable applications and improving software development processes."
}
work_experience_template = {
    "work_experience": [
        {
            "job_title": "Senior Software Engineer",
            "company": "Tech Solutions Inc.",
            "location": "San Francisco, CA",
            "date_range": "June 2018 - Present",
            "bullets": [
                "Led the development of a new feature that increased user engagement by 20%",
                "Implemented a CI/CD pipeline that reduced deployment time by 50%",
                "Mentored junior developers, improving team productivity and knowledge sharing",
            ],
        },
        {
            "job_title": "Software Engineer",
            "company": "Innovate Startup",
            "location": "San Francisco, CA",
            "date_range": "June 2016 - June 2018",
            "bullets": [
                "Developed a high-traffic web application using React and Node.js",
                "Optimized database queries, reducing load times by 30%",
                "Collaborated with cross-functional teams to define, design, and ship new features",
            ],
        },
    ],
}
additional_activities_template = (
    {
        "additional_activities": [
            {
                "title": "Hackathon Participant",
                "bullets": [
                    "Participated in a 24-hour hackathon, developing a web application that won 2nd place."
                ],
                "date": "March 2020",
            },
        ],
        "leadership": [
            {
                "title": "Director, Student Government Operations Council",
                "bullets": [
                    "Planned and executed over 15 events a term, communicating progress of the council weekly to other government sub-sections",
                    "Led a team of 20 students in planning and executing events",
                ],
                "date": "March 2020",
            },
        ],
    },
)
skills_template = {
    "skills": [
        "Python",
        "Java",
        "Docker",
        "Kubernetes",
        "React",
        "Agile methodologies",
    ]
}
education_template = {
    "education": [
        {
            "degree": "Bachelor of Science in Computer Science",
            "institution": "University of Example",
            "location": "Example City",
            "date_range": "September 2010 - June 2014",
            "bullets": "Graduated with honors",
        },
        {
            "degree": "Certified Kubernetes Administrator",
            "institution": "The Linux Foundation",
            "date_range": "July 2019",
        },
    ],
}
certifications_template = {
    "certifications": [
        {
            "title": "Certified Kubernetes Administrator",
            "issuer": "The Linux Foundation",
            "date_issued": "July 2019",
        }
    ],
}

# =========================== get resume sections =============================================================================================


def get_header_section(resume):
    start = time.time()
    resume = " ".join(resume.split()[:50])
    print(f"SPLIT resume: {resume}")
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": f"Please find the user header data, and return it in JSON format.",
            },
            {
                "role": "user",
                "content": f"Please find the user header data and return it in JSON format like this example: {header_template}. \n Here is the resume: {resume}.",
            },
        ],
        temperature=0.2,
    )
    section = completion.choices[0].message.content
    if section == "":
        print("No header found.")
        return None
    print(f"response: {section}")
    print(f"GPT time took: {time.time() - start} seconds.\n")
    return section


def get_summary_section(resume):
    start = time.time()
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": f"Please find the user summary statement data, and return it in JSON format.",
            },
            {
                "role": "user",
                "content": f"""Please find the user summary/objective statement and return it in JSON format like this example: {summary_template}. \n Here is the resume: {resume}. 
                If the user does not have a summary statement already present in their resume, please return an empty string.""",
            },
        ],
        temperature=0.2,
    )
    section = completion.choices[0].message.content
    if section == "":
        print("No summary found.")
        return None
    print(f"response: {section}")
    print(f"GPT time took: {time.time() - start} seconds.\n")
    return section


def get_work_experience_section(resume):
    start = time.time()
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": f"Please find the user work experiences from their resume, and return it in JSON format.",
            },
            {
                "role": "user",
                "content": f"""Please find the user work experiences from their resume and return it in JSON format like this example: {work_experience_template}. \n Here is the resume: {resume}. 
                If the user does not have an work experience already present in their resume, please return an empty string.""",
            },
        ],
        temperature=0.2,
    )
    section = completion.choices[0].message.content
    if section == "":
        print("No work experience found.")
        return None
    print(f"response for WE section: {section}")
    print(f"GPT time took: {time.time() - start} seconds.\n")
    return section


def get_additional_activities_section(resume):
    start = time.time()
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": f"Please find the user additional activities or miscellaneous section from their resume, and return it in JSON format.",
            },
            {
                "role": "user",
                "content": f"""Please find the user additional activities or miscellaneous section from their resume and return it in JSON format like this example: {additional_activities_template}. \n Here is the resume: {resume}. 
                If the user does not have an additional activities or miscellaneous section already present in their resume, please return an empty string.""",
            },
        ],
        temperature=0.2,
    )
    section = completion.choices[0].message.content
    if section == "":
        print("No additional activities found.")
        return None
    print(f"response for AA / Le section: {section}")
    print(f"GPT time took: {time.time() - start} seconds.\n")
    return section


def get_skills_section(resume):
    start = time.time()
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": f"Please find the user skills section from their resume, and return it in JSON format.",
            },
            {
                "role": "user",
                "content": f"""Please find the user skills section from their resume and return it in JSON format like this example: {json_template}. \n Here is the resume: {resume}. 
                If the user does not have a skills section already present in their resume, please return an empty string.""",
            },
        ],
        temperature=0.2,
    )
    section = completion.choices[0].message.content
    if section == "":
        print("No additional activities found.")
        return None
    print(f"response skills section: {section}")
    print(f"GPT time took: {time.time() - start} seconds.\n")
    return section


def get_education_section(resume):
    start = time.time()
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": f"Please find the user education section from their resume, and return it in JSON format.",
            },
            {
                "role": "user",
                "content": f"""Please find the user skills section from their resume and return it in JSON format like this example: {education_template}. \n Here is the resume: {resume}. 
                If the user does not have a education section already present in their resume, please return an empty string.""",
            },
        ],
        temperature=0.2,
    )
    section = completion.choices[0].message.content
    if section == "":
        print("No additional activities found.")
        return None
    print(f"response EDU: {section}")
    print(f"GPT time took: {time.time() - start} seconds.\n")
    return section


def get_certifications_section(resume):
    start = time.time()
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": f"Please find the user certifications section from their resume, and return it in JSON format.",
            },
            {
                "role": "user",
                "content": f"""Please find the user certifications section from their resume and return it in JSON format like this example: {certifications_template}. \n Here is the resume: {resume}. 
                If the user does not have a certifications section already present in their resume, please return an empty string.""",
            },
        ],
        temperature=0.2,
    )
    section = completion.choices[0].message.content
    if section == "":
        print("No additional activities found.")
        return None
    print(f"response: {section}")
    print(f"GPT time took: {time.time() - start} seconds.\n")
    return section


def convert_string_to_json(string):
    try:
        json_formatted = json.loads(string)
        return json_formatted
    except json.JSONDecodeError as e:
        print(f"Error: {e}")
        return None


def generate_sections(dummy_resume, concurrent=True):
    if concurrent == False:
        # ============================ Call the functions in serial to parse the sections ============================
        user_header_data = convert_string_to_json(get_header_section(dummy_resume))
        user_summary_data = convert_string_to_json(get_summary_section(dummy_resume))
        work_experience_data = convert_string_to_json(
            get_work_experience_section(dummy_resume)
        )
        additional_activities_data = convert_string_to_json(
            get_additional_activities_section(dummy_resume)
        )
        skills_data = convert_string_to_json(get_skills_section(dummy_resume))
        education_data = convert_string_to_json(get_education_section(dummy_resume))
        certifications_data = convert_string_to_json(
            get_certifications_section(dummy_resume)
        )
    else:
        # ============================ Call the functions concurrently to parse the sections ============================
        with ThreadPoolExecutor(max_workers=7) as executor:
            future_header = executor.submit(get_header_section, dummy_resume)
            future_summary = executor.submit(get_summary_section, dummy_resume)
            future_work_experience = executor.submit(
                get_work_experience_section, dummy_resume
            )
            future_additional_activities = executor.submit(
                get_additional_activities_section, dummy_resume
            )
            future_skills = executor.submit(get_skills_section, dummy_resume)
            future_education = executor.submit(get_education_section, dummy_resume)
            # future_certifications = executor.submit(get_certifications_section, dummy_resume)

            user_header_data = convert_string_to_json(future_header.result())
            user_summary_data = convert_string_to_json(future_summary.result())
            work_experience_data = convert_string_to_json(
                future_work_experience.result()
            )
            additional_activities_data = convert_string_to_json(
                future_additional_activities.result()
            )
            skills_data = convert_string_to_json(future_skills.result())
            education_data = convert_string_to_json(future_education.result())
            # certifications_data = convert_string_to_json(future_certifications.result())

    # combine the parsed sections into a single JSON object
    parsed_user_data = {
        **user_header_data,
        **user_summary_data,
        **work_experience_data,
        **additional_activities_data,
        **skills_data,
        **education_data,
        # **certifications_data,
    }
    return parsed_user_data


# =========================== Document manipulation functions ========================================================================================
def replace_text_while_keeping_formatting(paragraph, old_text, new_text):
    for run in paragraph.runs:
        if run.text.strip() != "":
            print(f"im changing this run: {run.text}")
            run.text = run.text.replace(old_text, new_text)
            print(f"BEFORE: {old_text} /// AFTER: {run.text}")


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        print(f"paragraph removed: {paragraph.text}")
        parent.remove(p)


def move_paragraph_after(paragraph, ref_paragraph):
    p = paragraph._element
    # print(f"paragraph text being moved: {paragraph.text}")
    ref_p = ref_paragraph._element
    parent = ref_p.getparent()
    if parent is not None:
        parent.insert(parent.index(ref_p) + 1, p)


def add_list_paragraph(document, paragraph, bullet_points, numId, ilvl=0):
    """
    Adds bullet points as a list to a Word document using OOXML.

    Parameters:
    - document: The Document object being modified.
    - paragraph: The Paragraph object where the list starts.
    - bullet_points: A list of strings, each representing a bullet point.
    - numId: The numbering definition ID to use for the list.
    - ilvl: The indentation level of the list (default is 0).
    """
    # Ensure the paragraph is part of a list
    p = paragraph._p  # access to the underlying lxml paragraph element
    pPr = p.get_or_add_pPr()  # access or add paragraph properties
    numPr = OxmlElement("w:numPr")  # create numPr element
    numId_element = OxmlElement("w:numId")
    numId_element.set(
        qn("w:val"), str(numId)
    )  # set numId to the numbering definition ID
    ilvl_element = OxmlElement("w:ilvl")
    ilvl_element.set(qn("w:val"), str(ilvl))  # set ilvl to the indentation level
    numPr.append(ilvl_element)
    numPr.append(numId_element)
    pPr.append(numPr)

    # Add bullet points as separate paragraphs
    for bullet in bullet_points:
        new_para = document.add_paragraph()
        new_para.add_run(bullet)
        # Apply the same list formatting
        new_p = new_para._p
        new_pPr = new_p.get_or_add_pPr()
        new_numPr = deepcopy(numPr)  # copy the numPr element to new paragraphs
        new_pPr.append(new_numPr)
        move_paragraph_after(new_para, paragraph)


# =========================== transfer data to template ========================================================================================


def add_work_experience_to_document(document, w_e_para_object, experiences):
    for experience in reversed(experiences):
        # Add Bullet Points
        if experience.get("bullets"):
            add_list_paragraph(document, w_e_para_object, experience["bullets"], numId=1)

        # Add Company Name and Location
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f"  {experience['company']}, {experience['location']} \t ({experience['date_range']})"
        )
        run.italic = True
        run.font.size = Pt(11)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, w_e_para_object)

        # Add Job Title
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(
            6
        )  # Add 6pt spacing before this paragraph
        run = paragraph.add_run(experience["job_title"])
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, w_e_para_object)

        # Add a blank line after each experience section for spacing
        document.add_paragraph()

    remove_paragraph(w_e_para_object)


def add_leadership_experience_to_document(document, l_e_para_object, experiences):
    for experience in reversed(experiences):
        # Add Bullet Points
        add_list_paragraph(document, l_e_para_object, experience["bullets"], numId=1)

        # Add leadership Title
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(
            6
        )  # Add 6pt spacing before this paragraph
        run = paragraph.add_run(experience["title"])
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, l_e_para_object)

    remove_paragraph(l_e_para_object)


def add_projects_to_document(document, p_para_object, projects):
    for project in reversed(projects):
        # Add Bullet Points
        add_list_paragraph(document, p_para_object, project["bullets"], numId=1)

        # Add Project Title
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(
            2
        )  # Add 6pt spacing before this paragraph
        run = paragraph.add_run(project["title"])
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, p_para_object)

    remove_paragraph(p_para_object)


def add_additional_activities_to_document(document, a_a_para_object, activities):
    for activity in reversed(activities):
        # Add Bullet Points
        if activity.get("bullets"):
            add_list_paragraph(document, a_a_para_object, activity["bullets"], numId=1)

        # Add Activity Title
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(
            2
        )  # Add 6pt spacing before this paragraph
        run = paragraph.add_run(activity["title"])
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, a_a_para_object)

    remove_paragraph(a_a_para_object)


def add_education_to_document(document, ed_para_object, educations):
    for education in reversed(educations):
        print(f"education: {education}")
        # Add Bullet Points
        if education.get("bullets"):
            add_list_paragraph(
                document, ed_para_object, education.get("bullets"), numId=1
            )

        # Add Company Name and Location
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f"  {education.get('institution')}, {education.get('location')} \t ({education.get('date_range')})"
        )
        run.italic = True
        run.font.size = Pt(11)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, ed_para_object)

        # Add Job Title
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(
            6
        )  # Add 6pt spacing before this paragraph
        run = paragraph.add_run(education["degree"])
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, ed_para_object)

    remove_paragraph(ed_para_object)


def combine_skills(skills):
    if isinstance(skills, dict):
        # Handle the case when skills are provided as a dictionary with categories
        new_text = ""
        for category, skills_list in skills.items():
            skills_text = ", ".join(skills_list)  # Join all skills in the list
            new_text += f"{category}: {skills_text}; "
    elif isinstance(skills, list):
        # Handle the case when skills are provided as a simple list
        new_text = ", ".join(skills)
    else:
        new_text = None

    return new_text


def transfer_data_to_template(para_object, document):

    # ==================================================== Header Section =========================================================
    if "Name" in para_object.text:
        new_text = parsed_user_data["personal_information"]["name"]
        replace_text_while_keeping_formatting(para_object, "Name", new_text)

    elif "information" in para_object.text:
        # TODO: Add a check for if the user has no linkedin, email, or phone
        components = [
            parsed_user_data["personal_information"].get("email"),
            parsed_user_data["personal_information"].get("phone"),
            parsed_user_data["personal_information"].get("linkedin"),
        ]
        # Filter out None values and join the remaining components with ' | '
        new_text = " | ".join(filter(None, components))
        replace_text_while_keeping_formatting(para_object, "information", new_text)

    # ==================================================== Summary Section =======================================================
    elif "summary_text" in para_object.text:
        new_text = parsed_user_data.get("summary")
        if new_text:
            print(f"found summary text")
            replace_text_while_keeping_formatting(para_object, "summary_text", new_text)
        else:
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if "Summary" in para_object.text:
                    remove_paragraph(para_object)

    # ==================================================== Additional Activities Section ======================================================
    elif "additional_activities_text" in para_object.text:
        additional_activities = parsed_user_data.get("additional_activities")
        if additional_activities:
            print(f"found additional_activities_text")
            add_additional_activities_to_document(
                document, para_object, additional_activities
            )
        else:
            # remove the placholder paragraph and the section header if the user has no activities section
            print(f"removing additional activities section")
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if ("additional_activities_text" in para_object.text) or (
                    "Additional Activities" in para_object.text
                ):
                    remove_paragraph(para_object)

    # ==================================================== Projects Section ======================================================
    elif "projects" in para_object.text:
        projects = parsed_user_data.get("projects")
        if projects:
            print(f"found work experience text in {para_object.text}")
            add_leadership_experience_to_document(document, para_object, projects)
        else:
            # remove the placholder paragraph and the section header if the user has no work experience
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if ("Projects" in para_object.text) or ("projects" in para_object.text):
                    remove_paragraph(para_object)

    # ==================================================== Leadership Section ====================================================
    elif "leadership_text" in para_object.text:
        leadership_experience = parsed_user_data.get("leadership_experience")
        if leadership_experience:
            print(f"found work experience text in {para_object.text}")
            add_leadership_experience_to_document(
                document, para_object, leadership_experience
            )
        else:
            # remove the placholder paragraph and the section header if the user has no work experience
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if ("Leadership Experience" in para_object.text) or (
                    "leadership_text" in para_object.text
                ):
                    remove_paragraph(para_object)

    # ==================================================== Work Experience Section ====================================================
    elif "work_experience_text" in para_object.text:
        work_experience = parsed_user_data.get("work_experience")
        if work_experience:
            print(f"found work experience text in {para_object.text}")
            add_work_experience_to_document(document, para_object, work_experience)
        else:
            # remove the placholder paragraph and the section header if the user has no work experience
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if ("Work Experience" in para_object.text) or (
                    "work_experience_text" in para_object.text
                ):
                    remove_paragraph(para_object)

    # ==================================================== Skills Section ========================================================
    elif "skills" in para_object.text:
        # Assuming summary text is a direct string in your new format
        skills_list = parsed_user_data.get("skills")
        new_text = combine_skills(skills_list)
        if new_text:
            print(f"found skill text")
            replace_text_while_keeping_formatting(para_object, "skills", new_text)
        else:
            # remove the placholder paragraph and the section header if the user has no skills
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if "Skills" in para_object.text:
                    remove_paragraph(para_object)

    # ==================================================== Education Section ====================================================
    elif "education_text" in para_object.text:
        if parsed_user_data.get("education"):
            print(f"found education text in {para_object.text}")
            add_education_to_document(
                document, para_object, parsed_user_data["education"]
            )
        else:
            # remove the placholder paragraph and the section header if the user has no education info
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if "Education" in para_object.text:
                    remove_paragraph(para_object)


# ===================================================== Function executions =================================================================

parsed_user_data = {
    "personal_information": {
        "name": "Cory G. Mazure",
        "email": "cory.mazure@gmail.com",
        "phone": "(586) 961-5274",
        "linkedin": "www.linkedin.com/in/cory-mazure",
    },
    "work_experience": [
        {
            "job_title": "Process Engineer, Golf Ball R&D",
            "company": "Callaway Golf",
            "location": "Carlsbad, CA",
            "date_range": "January 2023 - Present",
            "bullets": [
                "Lead engineer behind the experimentation, validation, and implementation of new spray gun equipment in production for 3 paint lines, resulting in a more controllable process and quality product",
                "Organized DOEs to study differences in mechanical processes between R&D and production, reducing variance in processes and accelerating new ball development",
                "Implemented new test methods to validate chemistry, leading to new proprietary golf ball features",
                "Utilized injection molding modeling software to modify tooling, improving concentricity and yield",
            ],
        },
        {
            "job_title": "Mechanical Engineering, Vehicle Safety Engineering",
            "company": "Stellantis",
            "location": "Auburn Hills, MI",
            "date_range": "Summer of 2019 – Summer 2022",
            "bullets": [
                "Led a design for six sigma project for optimizing the amount of test dummies in use, saving $250K/year",
                "Reduced vehicle validation timing by developing test fixtures for high inertia door latch testing",
                "Analyzed channel data for 80+ test rib deflection study, resulting in global program airbag innovations",
                "Coordinated JD Power quality study and benchmarking for heated steering wheel (HSW) optimization",
                "Strengthened best practices for HSW calibration ranges, resulting in enhanced customer satisfaction",
                "Investigated 2 competing suppliers for hands-on detection mats via Minitab statistical analysis, leading to more competitive pricing",
                "Established statistical analysis of the dummy labs 4 certification machines, reducing certification timing and associated costs by $20K/year",
                "Validated 2 energy absorbing knee brackets and a driver airbag, achieving higher safety ratings",
            ],
        },
    ],
    "leadership_experience": [
        {
            "title": "Kettering Student Government Operations Council – Director",
            "bullets": [
                "Planned and executed over 15 events a term, communicating progress of the council weekly to other government sub-sections"
            ],
        },
    ],
    "projects": [
        {
            "title": "Personal Portfolio Website",
            "bullets": [
                "A personal website to showcase my projects and resume. Built with HTML, CSS, and JavaScript."
            ],
            "link": "http://www.johndoeportfolio.com",
        }
    ],
    "education": [
        {
            "degree": "Bachelor of Science in Mechanical Engineering",
            "institution": "Kettering University",
            "location": "Flint, MI",
            "date_range": "Oct. 2019 - Dec. 2022",
            "bullets": [
                "Graduated with honors",
                "GPA: 3.95/4.00",
                "Summa Cum Laude",
                "Dean's List every academic term",
            ],
        },
        {
            "degree": "Kubernetes Administrator",
            "institution": "The International Linux Foundation",
            "date_range": "Oct. 2019 - Dec. 2022",
        },
    ],
    "skills": [
        "Stellantis Sponsored Courses: Attended 5 in depth training courses for DFMEA, DFSS, and GD&T",
        "Moldflow injection molding simulation software",
        "Proficient in NX, Fusion 360, & SolidWorks: CAE and FEA",
        "Experienced in Minitab statistical analysis software for interpreting DOE results",
        "Utilized Altair Hyperworks to analyze vehicle crash data",
        "Avid user of Microsoft Office and Google suite for statistical analysis and project management",
    ],
}


generation_start = time.time()
parsed_user_data = generate_sections(resume_text)
json_user_data = json.dumps(parsed_user_data, indent=4)
# write to json file for review:
with open(
    "/Users/campbmaso/Desktop/Development/GitHub/Sandbox/Backend/resumes/parsed_user_data.json",
    "w",
) as file:
    file.write(json_user_data)
    print(f"parsed_user_data written to file.")

print(f"\n\nparsed_user_data: {parsed_user_data}\n\n")

print(f"Total generation_time took: {time.time() - generation_start} seconds.\n")


loop_start = time.time()

processed_sections = set()
for para_object in list(new_doc.paragraphs):
    if para_object.text not in processed_sections:
        # print(f"para_object.text: {para_object.text}")
        transfer_data_to_template(para_object, new_doc)
        processed_sections.add(para_object.text)
    else:
        print(f"Skipping already processed para_object.text: {para_object.text}")

# print(f"processed_sections: {processed_sections}")

# ===============================================================================================================================================
# =========================== save the new document =============================================================================================
# ===============================================================================================================================================

# Save the modified document back to a BytesIO stream
modified_content = BytesIO()
new_doc.save(modified_content)

# Now, to save the content of the BytesIO object to a file
with open(
    "/Users/campbmaso/Desktop/Development/GitHub/Sandbox/Backend/resumes/RESULT_resume.docx",
    "wb",
) as output_file:
    # Go to the beginning of the BytesIO stream
    modified_content.seek(0)
    # Write the contents of the BytesIO stream to the file
    output_file.write(modified_content.read())
    print("File saved successfully!")

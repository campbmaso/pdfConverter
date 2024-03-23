import os
import json
import time
from io import BytesIO

import boto3
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy

s3_client = boto3.client('s3')


# =========================== Document manipulation functions ========================================================================================
def replace_text_while_keeping_formatting(paragraph, old_text, new_text):
    for run in paragraph.runs:
        if run.text.strip() != "":
            print(f"im changing this run: {run.text}")
            run.text = run.text.replace(old_text, new_text)
            print(f"BEFORE: {old_text} /// AFTER: {run.text}")

def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        print(f"paragraph removed: {paragraph.text}")
        parent.remove(p)

def move_paragraph_after(paragraph, ref_paragraph):
    p = paragraph._element
    # print(f"paragraph text being moved: {paragraph.text}")
    ref_p = ref_paragraph._element
    parent = ref_p.getparent()
    if parent is not None:
        parent.insert(parent.index(ref_p) + 1, p)

def add_list_paragraph(document, paragraph, bullet_points, numId, ilvl=0):
    """
    Adds bullet points as a list to a Word document using OOXML.

    Parameters:
    - document: The Document object being modified.
    - paragraph: The Paragraph object where the list starts.
    - bullet_points: A list of strings, each representing a bullet point.
    - numId: The numbering definition ID to use for the list.
    - ilvl: The indentation level of the list (default is 0).
    """
    # Ensure the paragraph is part of a list
    p = paragraph._p  # access to the underlying lxml paragraph element
    pPr = p.get_or_add_pPr()  # access or add paragraph properties
    numPr = OxmlElement("w:numPr")  # create numPr element
    numId_element = OxmlElement("w:numId")
    numId_element.set(
        qn("w:val"), str(numId)
    )  # set numId to the numbering definition ID
    ilvl_element = OxmlElement("w:ilvl")
    ilvl_element.set(qn("w:val"), str(ilvl))  # set ilvl to the indentation level
    numPr.append(ilvl_element)
    numPr.append(numId_element)
    pPr.append(numPr)

    # Ensure bullet_points is a list
    if isinstance(bullet_points, str):
        bullet_points = [bullet_points]  # Convert a single string into a list
        
    # Add bullet points as separate paragraphs
    for bullet in bullet_points:
        new_para = document.add_paragraph()
        new_para.add_run(bullet)
        # Apply the same list formatting
        new_p = new_para._p
        new_pPr = new_p.get_or_add_pPr()
        new_numPr = deepcopy(numPr)  # copy the numPr element to new paragraphs
        new_pPr.append(new_numPr)
        move_paragraph_after(new_para, paragraph)

# =========================== transfer data to template ========================================================================================

def add_work_experience_to_document(document, w_e_para_object, experiences):
    for experience in reversed(experiences):
        # Add Bullet Points
        if experience.get("bullets"):
            add_list_paragraph(document, w_e_para_object, experience["bullets"], numId=1)

        # Add Company Name and Location
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f"  {experience['company']}, {experience['location']} \t ({experience['date_range']})"
        )
        run.italic = True
        run.font.size = Pt(11)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, w_e_para_object)

        # Add Job Title
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(
            6
        )  # Add 6pt spacing before this paragraph
        run = paragraph.add_run(experience["job_title"])
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, w_e_para_object)

        # Add a blank line after each experience section for spacing
        document.add_paragraph()

    remove_paragraph(w_e_para_object)

def add_leadership_experience_to_document(document, l_e_para_object, experiences):
    for experience in reversed(experiences):
        # Add Bullet Points
        add_list_paragraph(document, l_e_para_object, experience["bullets"], numId=1)

        # Add leadership Title
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(
            6
        )  # Add 6pt spacing before this paragraph
        run = paragraph.add_run(experience["title"])
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, l_e_para_object)

    remove_paragraph(l_e_para_object)

def add_projects_to_document(document, p_para_object, projects):
    for project in reversed(projects):
        # Add Bullet Points
        add_list_paragraph(document, p_para_object, project["bullets"], numId=1)

        # Add Project Title
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(
            2
        )  # Add 6pt spacing before this paragraph
        run = paragraph.add_run(project["title"])
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, p_para_object)

    remove_paragraph(p_para_object)

def add_additional_activities_to_document(document, a_a_para_object, activities):
    for activity in reversed(activities):
        # Add Bullet Points
        if activity.get("bullets"):
            add_list_paragraph(document, a_a_para_object, activity["bullets"], numId=1)

        # Add Activity Title
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(
            2
        )  # Add 6pt spacing before this paragraph
        run = paragraph.add_run(activity["title"])
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, a_a_para_object)

    remove_paragraph(a_a_para_object)

def add_education_to_document(document, ed_para_object, educations):
    for education in reversed(educations):
        print(f"education: {education}")
        # Add Bullet Points
        if education.get("bullets"):
            add_list_paragraph(
                document, ed_para_object, education.get("bullets"), numId=1
            )

        # Add Company Name and Location
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            f"  {education.get('institution')}, {education.get('location')} \t ({education.get('date_range')})"
        )
        run.italic = True
        run.font.size = Pt(11)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, ed_para_object)

        # Add Job Title
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(
            6
        )  # Add 6pt spacing before this paragraph
        run = paragraph.add_run(education["degree"])
        run.bold = True
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        move_paragraph_after(paragraph, ed_para_object)

    remove_paragraph(ed_para_object)

def combine_skills(skills):
    if isinstance(skills, dict):
        # Handle the case when skills are provided as a dictionary with categories
        new_text = ""
        for category, skills_list in skills.items():
            skills_text = ", ".join(skills_list)  # Join all skills in the list
            new_text += f"{category}: {skills_text}; "
    elif isinstance(skills, list):
        # Handle the case when skills are provided as a simple list
        new_text = ", ".join(skills)
    else:
        new_text = None

    return new_text

def transfer_data_to_template(para_object, new_doc, parsed_user_data):
    
    # ==================================================== Header Section =========================================================
    if "Name" in para_object.text:
        new_text = parsed_user_data["personal_information"]["name"]
        replace_text_while_keeping_formatting(para_object, "Name", new_text)

    elif "information" in para_object.text:
        # TODO: Add a check for if the user has no linkedin, email, or phone
        components = [
            parsed_user_data["personal_information"].get("email"),
            parsed_user_data["personal_information"].get("phone"),
            parsed_user_data["personal_information"].get("linkedin"),
        ]
        # Filter out None values and join the remaining components with ' | '
        new_text = " | ".join(filter(None, components))
        replace_text_while_keeping_formatting(para_object, "information", new_text)

    # ==================================================== Summary Section =======================================================
    elif "summary_text" in para_object.text:
        new_text = parsed_user_data.get("summary")
        if new_text:
            print(f"found summary text")
            replace_text_while_keeping_formatting(para_object, "summary_text", new_text)
        else:
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if "Summary" in para_object.text:
                    remove_paragraph(para_object)

    # ==================================================== Additional Activities Section ======================================================
    elif "additional_activities_text" in para_object.text:
        additional_activities = parsed_user_data.get("additional_activities")
        if additional_activities:
            print(f"found additional_activities_text")
            add_additional_activities_to_document(
                new_doc, para_object, additional_activities
            )
        else:
            # remove the placholder paragraph and the section header if the user has no activities section
            print(f"removing additional activities section")
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if ("additional_activities_text" in para_object.text) or (
                    "Additional Activities" in para_object.text
                ):
                    remove_paragraph(para_object)

    # ==================================================== Projects Section ======================================================
    elif "projects" in para_object.text:
        projects = parsed_user_data.get("projects")
        if projects:
            print(f"found work experience text in {para_object.text}")
            add_leadership_experience_to_document(new_doc, para_object, projects)
        else:
            # remove the placholder paragraph and the section header if the user has no work experience
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if ("Projects" in para_object.text) or ("projects" in para_object.text):
                    remove_paragraph(para_object)

    # ==================================================== Leadership Section ====================================================
    elif "leadership_text" in para_object.text:
        leadership_experience = parsed_user_data.get("leadership_experience")
        if leadership_experience:
            print(f"found work experience text in {para_object.text}")
            add_leadership_experience_to_document(
                new_doc, para_object, leadership_experience
            )
        else:
            # remove the placholder paragraph and the section header if the user has no work experience
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if ("Leadership Experience" in para_object.text) or (
                    "leadership_text" in para_object.text
                ):
                    remove_paragraph(para_object)

    # ==================================================== Work Experience Section ====================================================
    elif "work_experience_text" in para_object.text:
        work_experience = parsed_user_data.get("work_experience")
        if work_experience:
            print(f"found work experience text in {para_object.text}")
            add_work_experience_to_document(new_doc, para_object, work_experience)
        else:
            # remove the placholder paragraph and the section header if the user has no work experience
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if ("Work Experience" in para_object.text) or ("work_experience_text" in para_object.text):
                    remove_paragraph(para_object)

    # ==================================================== Skills Section ========================================================
    elif "skills" in para_object.text:
        # Assuming summary text is a direct string in your new format
        skills_list = parsed_user_data.get("skills")
        new_text = combine_skills(skills_list)
        if new_text:
            print(f"found skill text")
            replace_text_while_keeping_formatting(para_object, "skills", new_text)
        else:
            # remove the placholder paragraph and the section header if the user has no skills
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if "Skills" in para_object.text:
                    remove_paragraph(para_object)

    # ==================================================== Education Section ====================================================
    elif "education_text" in para_object.text:
        if parsed_user_data.get("education"):
            print(f"found education text in {para_object.text}")
            add_education_to_document(
                new_doc, para_object, parsed_user_data["education"]
            )
        else:
            # remove the placholder paragraph and the section header if the user has no education info
            remove_paragraph(para_object)
            for para_object in new_doc.paragraphs:
                if "Education" in para_object.text:
                    remove_paragraph(para_object)

def convert_pdf_to_docx2(file_content, filename):
    start = time.time()
    data = file_content
    if not data.get('errorMessage'):
        parsed_user_data = json.loads(data.get('body'))
        
        # grabbing the template resume from s3
        template_version = 1
        template_object = s3_client.get_object(
            Bucket="resume-s3bucket", Key=f"templates/Serif Template {template_version}.docx"
        )
        template_content = BytesIO(template_object["Body"].read())

        new_doc = Document(template_content)
        print("printing new_doc paragraphs BEFORE MODS")
        for para_object in list(new_doc.paragraphs):
            print(f"para_object.text: {para_object.text}")

        # transferring data to the template    
        processed_sections = set()
        for para_object in list(new_doc.paragraphs):
            if para_object.text not in processed_sections:
                # print(f"para_object.text: {para_object.text}")
                transfer_data_to_template(para_object, new_doc, parsed_user_data)
                processed_sections.add(para_object.text)
            else:
                print(f"Skipping already processed para_object.text: {para_object.text}")
        # Save the modified document back to a BytesIO stream
        modified_content = BytesIO()
        new_doc.save(modified_content)
        # Reset the stream's position to the beginning
        modified_content.seek(0)
        print("printing new_doc paragraphs after MODS")
        for para_object in list(new_doc.paragraphs):
            print(f"para_object.text AFTER MODS: {para_object.text}")
        # resume_text = "okay"
        return modified_content